#!/usr/bin/env python3
"""
extract_bidders.py — Extract bidder data from opening protocol documents.

Scans all tender folders, identifies the best opening protocol document(s),
extracts bidder tables (company, lot, bid amount, timestamp), and writes
structured output to Excel.

Document priority:
  1. PROPFIS docx (financial opening — has bid amounts)
  2. OPNPRT docx  (participation opening — has bidders + timestamps)
  3. Any docx with protocol/atversana keywords
  4. PDF protocols (pdfplumber, then OpenAI vision fallback)

Usage:
  python extract_bidders.py --downloads ../downloads --inventory ../inventory.xlsx
  python extract_bidders.py --downloads ../downloads --inventory ../inventory.xlsx --limit 10
  python extract_bidders.py --downloads ../downloads --inventory ../inventory.xlsx --procurement-ids 156048,108572
"""

import argparse
import json
import os
import re
import sys
import time
import unicodedata
from datetime import datetime
from pathlib import Path

import pandas as pd
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

# ═══════════════════════════════════════════════════════════════════════════════
# CONFIGURATION
# ═══════════════════════════════════════════════════════════════════════════════
OPENAI_API_KEY = os.environ.get("OPENAI_API_KEY", "")
OPENAI_MODEL = os.environ.get("OPENAI_MODEL", "gpt-4.1-mini")
MIN_SECONDS_BETWEEN_CALLS = 1.0

SUPPORTED_EXT = {".docx", ".doc", ".pdf"}
SKIP_PATTERNS = [
    "document-cert", "mimetype", "manifest.xml", "signatures",
    ".xml", ".7z", ".zip", ".edoc", "pd.any", "meta-inf",
]
NON_PROTOCOL_KEYWORDS = [
    "nolikums", "nolikuma", "ligums", "vienosanas", "contract",
    "specifikacija", "pieteikums", "pieteikuma", "zinojums",
    "lemums", "atbilde", "grozij", "preciz",
]

# ═══════════════════════════════════════════════════════════════════════════════
# TEXT NORMALISATION
# ═══════════════════════════════════════════════════════════════════════════════
def normalise(text: str) -> str:
    if not text:
        return ""
    s = str(text).lower()
    s = unicodedata.normalize("NFKD", s)
    s = "".join(ch for ch in s if not unicodedata.combining(ch))
    s = re.sub(r"[_\-\./\\|:;,\(\)\[\]\{\}\"]+", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


# ═══════════════════════════════════════════════════════════════════════════════
# DOCUMENT SELECTION — find the best protocol file(s) in a tender folder
# ═══════════════════════════════════════════════════════════════════════════════
def _is_skip(fname: str) -> bool:
    fl = fname.lower()
    return any(p in fl for p in SKIP_PATTERNS)


def _score_protocol_file(fname: str) -> tuple:
    """
    Returns (priority, is_financial) where lower priority = better.
    Priority 0 = PROPFIS docx (financial opening, best)
    Priority 1 = OPNPRT docx (participation opening)
    Priority 2 = other protocol-keyword docx
    Priority 3 = PROPFIS/OPNPRT pdf
    Priority 4 = other protocol pdf
    Priority 99 = not a protocol
    """
    fl = fname.lower()
    norm = normalise(fname)
    ext = Path(fname).suffix.lower()

    if _is_skip(fname) or ext not in SUPPORTED_EXT:
        return (99, False)

    # Exclude files that are clearly other document types
    if re.search(r"\d+\s+pielikums\b", norm):
        return (99, False)
    for kw in NON_PROTOCOL_KEYWORDS:
        if kw in norm and "protokol" not in norm and "opnprt" not in norm and "openingprotocol" not in norm:
            return (99, False)

    is_financial = "propfis" in fl
    is_opnprt = "opnprt" in fl and not is_financial
    is_opening = "openingprotocol" in fl or "meeting_openingprotocol" in fl
    has_proto_kw = any(kw in norm for kw in [
        "protokols", "protokol", "opnprt", "atversana", "atversan",
        "opening protocol", "meeting openingprotocol",
    ])

    if ext == ".docx" or ext == ".doc":
        if is_financial:
            return (0, True)
        if is_opnprt or is_opening:
            return (1, False)
        if has_proto_kw:
            return (2, False)
    elif ext == ".pdf":
        if is_financial or is_opnprt or is_opening:
            return (3, is_financial)
        if has_proto_kw:
            return (4, False)

    return (99, False)


def find_protocol_files(folder: Path) -> list[dict]:
    """Find and rank all protocol files in a tender folder."""
    if not folder.is_dir():
        return []

    candidates = []
    seen_norm = set()
    for f in sorted(folder.iterdir()):
        if not f.is_file():
            continue
        priority, is_fin = _score_protocol_file(f.name)
        if priority >= 99:
            continue
        # Dedup by normalised name (strip ID prefix)
        norm = re.sub(r"^id\d+_", "", normalise(f.name))
        if norm in seen_norm:
            continue
        seen_norm.add(norm)
        candidates.append({
            "path": f,
            "filename": f.name,
            "priority": priority,
            "is_financial": is_fin,
        })

    candidates.sort(key=lambda x: x["priority"])
    return candidates


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX TABLE EXTRACTION
# ═══════════════════════════════════════════════════════════════════════════════
def _parse_amount(text: str) -> float | None:
    """Parse EUR amount from cell text like 'EUR 14950.0' or '14 950,00'."""
    if not text:
        return None
    s = text.strip()
    s = re.sub(r"(?i)^eur\s*", "", s)
    s = s.replace("\xa0", "").replace(" ", "")
    # Handle comma as decimal separator
    if "," in s and "." in s:
        s = s.replace(",", "")
    elif "," in s:
        s = s.replace(",", ".")
    try:
        return float(s)
    except ValueError:
        return None


def _parse_timestamp(text: str) -> str:
    """Normalise timestamp like '16.10.2025 plkst. 18:02'."""
    if not text:
        return ""
    s = text.strip()
    s = re.sub(r"\s*plkst\.\s*", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s


def _identify_column_roles(header_cells: list[str]) -> dict:
    """Map column indices to roles based on header text."""
    roles = {}
    for i, h in enumerate(header_cells):
        hn = normalise(h)
        if not hn:
            continue
        if "pretendents" in hn or "bidder" in hn:
            roles["bidder"] = i
        elif "cena" in hn or "summa" in hn or "finansu" in hn or "piedavajum" in hn or "price" in hn:
            if "amount" not in roles:
                roles["amount"] = i
        elif "datums" in hn or "laiks" in hn or "iesniegsanas" in hn:
            if "timestamp" not in roles:
                roles["timestamp"] = i
        elif "n.p.k" in hn or "nr." in hn:
            roles["npk"] = i
    return roles


def _extract_lot_number(text: str) -> int | None:
    """Extract lot number from text like 'Daļai Nr. 3 - ...'."""
    m = re.search(r"(?:Da[lļ]ai|Lot)\s*Nr\.?\s*(\d+)", text, re.IGNORECASE)
    return int(m.group(1)) if m else None


def _is_data_row(cells: list[str], roles: dict) -> bool:
    """Check if a row contains actual bidder data (not header/empty)."""
    bidder_idx = roles.get("bidder")
    if bidder_idx is None:
        return False
    bidder = cells[bidder_idx].strip() if bidder_idx < len(cells) else ""
    if not bidder:
        return False
    bn = normalise(bidder)
    if any(kw in bn for kw in ["pretendents", "bidder", "n.p.k", "iesniegsanas", "datums"]):
        return False
    if len(bidder) < 3:
        return False
    return True


def extract_from_docx(filepath: Path, procurement_id: int) -> list[dict]:
    """Extract bidder records from a docx protocol file."""
    from docx import Document
    try:
        doc = Document(str(filepath))
    except Exception as e:
        return [{"procurement_id": procurement_id, "error": f"docx_read_error: {e}",
                 "source_file": filepath.name}]

    if not doc.tables:
        return [{"procurement_id": procurement_id, "error": "no_tables_found",
                 "source_file": filepath.name}]

    records = []
    current_lot = None
    is_financial = "propfis" in filepath.name.lower()

    for table in doc.tables:
        rows_data = []
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            rows_data.append(cells)

        if not rows_data:
            continue

        # Check for lot header in first row
        first_text = " ".join(rows_data[0])
        lot_num = _extract_lot_number(first_text)
        if lot_num is not None:
            current_lot = lot_num

        # Find the header row with "Pretendents"
        header_idx = None
        for i, cells in enumerate(rows_data):
            joined = normalise(" ".join(cells))
            if "pretendents" in joined:
                header_idx = i
                break

        if header_idx is None:
            continue

        roles = _identify_column_roles(rows_data[header_idx])
        if "bidder" not in roles:
            continue

        # Extract data rows
        for cells in rows_data[header_idx + 1:]:
            if not _is_data_row(cells, roles):
                continue

            bidder_idx = roles["bidder"]
            bidder = cells[bidder_idx].strip() if bidder_idx < len(cells) else ""

            amount = None
            if "amount" in roles and roles["amount"] < len(cells):
                amount = _parse_amount(cells[roles["amount"]])

            timestamp = ""
            if "timestamp" in roles and roles["timestamp"] < len(cells):
                timestamp = _parse_timestamp(cells[roles["timestamp"]])

            records.append({
                "procurement_id": procurement_id,
                "lot": current_lot if current_lot else 1,
                "bidder_name": bidder,
                "bid_amount": amount,
                "bid_timestamp": timestamp,
                "source_file": filepath.name,
                "source_type": "financial" if is_financial else "participation",
                "confidence": "high",
                "error": "",
            })

    # Deduplicate within same file (e.g. both participation + financial table)
    records = _merge_records(records)
    return records


# ═══════════════════════════════════════════════════════════════════════════════
# PDF TABLE EXTRACTION
# ═══════════════════════════════════════════════════════════════════════════════
def extract_from_pdf_pdfplumber(filepath: Path, procurement_id: int) -> list[dict]:
    """Try extracting bidder tables from PDF using pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        return []

    records = []
    current_lot = None

    try:
        with pdfplumber.open(str(filepath)) as pdf:
            for page in pdf.pages:
                # Check page text for lot indicators
                page_text = page.extract_text() or ""
                lot_num = _extract_lot_number(page_text)
                if lot_num is not None:
                    current_lot = lot_num

                tables = page.extract_tables()
                if not tables:
                    continue

                for table in tables:
                    if not table or len(table) < 2:
                        continue

                    # Find header row
                    header_idx = None
                    for i, row in enumerate(table):
                        joined = normalise(" ".join(str(c or "") for c in row))
                        if "pretendents" in joined:
                            header_idx = i
                            break

                    if header_idx is None:
                        continue

                    header_cells = [str(c or "") for c in table[header_idx]]
                    roles = _identify_column_roles(header_cells)
                    if "bidder" not in roles:
                        continue

                    for row in table[header_idx + 1:]:
                        cells = [str(c or "").strip() for c in row]
                        if not _is_data_row(cells, roles):
                            continue

                        bidder = cells[roles["bidder"]] if roles["bidder"] < len(cells) else ""
                        amount = None
                        if "amount" in roles and roles["amount"] < len(cells):
                            amount = _parse_amount(cells[roles["amount"]])
                        timestamp = ""
                        if "timestamp" in roles and roles["timestamp"] < len(cells):
                            timestamp = _parse_timestamp(cells[roles["timestamp"]])

                        records.append({
                            "procurement_id": procurement_id,
                            "lot": current_lot if current_lot else 1,
                            "bidder_name": bidder,
                            "bid_amount": amount,
                            "bid_timestamp": timestamp,
                            "source_file": filepath.name,
                            "source_type": "pdf_pdfplumber",
                            "confidence": "medium",
                            "error": "",
                        })
    except Exception as e:
        return [{"procurement_id": procurement_id, "error": f"pdfplumber_error: {e}",
                 "source_file": filepath.name}]

    return records


def extract_from_pdf_llm(filepath: Path, procurement_id: int) -> list[dict]:
    """Fallback: send PDF pages as images to OpenAI for table extraction."""
    if not OPENAI_API_KEY:
        return [{"procurement_id": procurement_id, "error": "missing_openai_api_key",
                 "source_file": filepath.name}]

    try:
        from openai import OpenAI
        import fitz  # PyMuPDF
        import base64
    except ImportError as e:
        return [{"procurement_id": procurement_id, "error": f"missing_lib: {e}",
                 "source_file": filepath.name}]

    try:
        doc = fitz.open(str(filepath))
    except Exception as e:
        return [{"procurement_id": procurement_id, "error": f"pdf_open_error: {e}",
                 "source_file": filepath.name}]

    # Render up to 6 pages as images
    images = []
    for page_num in range(min(len(doc), 6)):
        page = doc[page_num]
        pix = page.get_pixmap(dpi=150)
        img_bytes = pix.tobytes("png")
        b64 = base64.b64encode(img_bytes).decode("utf-8")
        images.append({"type": "image_url", "image_url": {"url": f"data:image/png;base64,{b64}"}})
    doc.close()

    prompt = """Extract ALL bidder data from this procurement opening protocol document.
Return a JSON array of objects with these fields:
- lot: integer lot number (1 if single lot/no lots mentioned)
- bidder_name: company name exactly as written
- bid_amount: numeric value only (no currency), null if not shown
- bid_timestamp: submission date/time as written, empty string if not shown

Return ONLY the JSON array, no other text. Example:
[{"lot":1,"bidder_name":"SIA Example","bid_amount":14950.0,"bid_timestamp":"16.10.2025 18:02"}]"""

    client = OpenAI(api_key=OPENAI_API_KEY)
    try:
        response = client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[{
                "role": "user",
                "content": [{"type": "text", "text": prompt}] + images,
            }],
            max_tokens=4096,
            temperature=0,
        )
        text = response.choices[0].message.content.strip()
        # Extract JSON from response
        text = re.sub(r"^```json\s*", "", text)
        text = re.sub(r"\s*```$", "", text)
        data = json.loads(text)

        records = []
        for item in data:
            records.append({
                "procurement_id": procurement_id,
                "lot": item.get("lot", 1),
                "bidder_name": item.get("bidder_name", ""),
                "bid_amount": item.get("bid_amount"),
                "bid_timestamp": item.get("bid_timestamp", ""),
                "source_file": filepath.name,
                "source_type": "pdf_llm",
                "confidence": "medium",
                "error": "",
            })
        return records

    except Exception as e:
        return [{"procurement_id": procurement_id, "error": f"llm_error: {e}",
                 "source_file": filepath.name}]


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN EXTRACTION PIPELINE
# ═══════════════════════════════════════════════════════════════════════════════
def extract_tender(folder: Path, procurement_id: int, skip_llm: bool = False) -> list[dict]:
    """Extract bidder data from all relevant protocol files for one tender."""
    protocol_files = find_protocol_files(folder)

    if not protocol_files:
        return [{
            "procurement_id": procurement_id,
            "lot": None, "bidder_name": "", "bid_amount": None,
            "bid_timestamp": "", "source_file": "", "source_type": "",
            "confidence": "", "error": "no_protocol_found",
        }]

    all_records = []
    used_sources = set()

    # Process files in priority order; stop after getting good data
    for pf in protocol_files:
        fp = pf["path"]
        ext = fp.suffix.lower()

        if ext in (".docx", ".doc"):
            records = extract_from_docx(fp, procurement_id)
        elif ext == ".pdf":
            records = extract_from_pdf_pdfplumber(fp, procurement_id)
            # If pdfplumber got nothing useful, try LLM
            valid = [r for r in records if not r.get("error") and r.get("bidder_name")]
            if not valid and not skip_llm:
                records = extract_from_pdf_llm(fp, procurement_id)
        else:
            continue

        valid_records = [r for r in records if not r.get("error") and r.get("bidder_name")]
        error_records = [r for r in records if r.get("error")]

        if valid_records:
            all_records.extend(valid_records)
            used_sources.add(fp.name)
            # If we got a financial doc (with amounts), that's best — check if we
            # need the participation doc too for completeness
            if pf["is_financial"]:
                break
        elif error_records:
            all_records.extend(error_records)

    # Merge: if we have both financial and participation records for same
    # bidder+lot, prefer the financial record (has amount)
    if len(used_sources) > 1:
        all_records = _merge_records(all_records)

    return all_records


def _merge_records(records: list[dict]) -> list[dict]:
    """Deduplicate records when multiple source files cover the same bidders."""
    # Group by (lot, normalised bidder name)
    by_key = {}
    for r in records:
        if r.get("error"):
            continue
        key = (r.get("lot", 1), normalise(r.get("bidder_name", "")))
        existing = by_key.get(key)
        if existing is None:
            by_key[key] = r
        else:
            # Prefer the one with bid_amount
            if r.get("bid_amount") is not None and existing.get("bid_amount") is None:
                by_key[key] = r
            # Merge timestamp if missing
            if not existing.get("bid_timestamp") and r.get("bid_timestamp"):
                existing["bid_timestamp"] = r["bid_timestamp"]

    result = list(by_key.values())
    # Add back error records
    result.extend(r for r in records if r.get("error"))
    return result


def _validate_records(records: list[dict]) -> list[dict]:
    """Flag low-confidence records."""
    for r in records:
        if r.get("error"):
            r["confidence"] = "error"
            continue

        issues = []
        if not r.get("bidder_name"):
            issues.append("missing_bidder")
        if r.get("bid_amount") is not None and r["bid_amount"] <= 0:
            issues.append("zero_or_negative_amount")
        if r.get("bid_amount") is not None and r["bid_amount"] > 1e9:
            issues.append("suspiciously_large_amount")

        if issues:
            r["confidence"] = "low"
            r["error"] = "; ".join(issues)

    return records


# ═══════════════════════════════════════════════════════════════════════════════
# OUTPUT
# ═══════════════════════════════════════════════════════════════════════════════
def write_output(all_records: list[dict], output_path: Path, tender_stats: dict):
    """Write extraction results to Excel with bid_level + summary sheets."""
    wb = openpyxl.Workbook()

    hdr_fill = PatternFill("solid", fgColor="4472C4")
    hdr_font = Font(name="Arial", bold=True, size=10, color="FFFFFF")
    normal_font = Font(name="Arial", size=10)
    err_fill = PatternFill("solid", fgColor="F4B084")
    low_fill = PatternFill("solid", fgColor="FFEB9C")
    ok_fill = PatternFill("solid", fgColor="C6EFCE")
    thin = Border(left=Side("thin"), right=Side("thin"), top=Side("thin"), bottom=Side("thin"))

    # --- Sheet 1: bid_level ---
    ws = wb.active
    ws.title = "bid_level"
    cols = ["procurement_id", "lot", "bidder_name", "bid_amount",
            "bid_timestamp", "source_file", "source_type", "confidence", "error"]
    for j, c in enumerate(cols, 1):
        cell = ws.cell(row=1, column=j, value=c)
        cell.font, cell.fill, cell.border = hdr_font, hdr_fill, thin

    for i, r in enumerate(all_records, 2):
        conf = r.get("confidence", "")
        fill = err_fill if conf == "error" else (low_fill if conf == "low" else ok_fill)
        for j, c in enumerate(cols, 1):
            cell = ws.cell(row=i, column=j, value=r.get(c, ""))
            cell.font, cell.border, cell.fill = normal_font, thin, fill

    widths = [14, 6, 40, 14, 22, 55, 16, 10, 30]
    for j, w in enumerate(widths, 1):
        ws.column_dimensions[chr(64 + j)].width = w
    ws.auto_filter.ref = f"A1:{chr(64+len(cols))}{len(all_records)+1}"

    # --- Sheet 2: tender_summary ---
    ws2 = wb.create_sheet("tender_summary")
    sum_cols = ["procurement_id", "n_bidders", "n_lots", "has_amounts",
                "protocol_type", "status"]
    for j, c in enumerate(sum_cols, 1):
        cell = ws2.cell(row=1, column=j, value=c)
        cell.font, cell.fill, cell.border = hdr_font, hdr_fill, thin

    for i, (pid, stats) in enumerate(sorted(tender_stats.items()), 2):
        vals = [pid, stats["n_bidders"], stats["n_lots"], stats["has_amounts"],
                stats["protocol_type"], stats["status"]]
        fill = err_fill if stats["status"] == "error" else (
            low_fill if stats["status"] == "low_confidence" else ok_fill)
        for j, v in enumerate(vals, 1):
            cell = ws2.cell(row=i, column=j, value=v)
            cell.font, cell.border, cell.fill = normal_font, thin, fill

    widths2 = [14, 12, 8, 14, 18, 16]
    for j, w in enumerate(widths2, 1):
        ws2.column_dimensions[chr(64 + j)].width = w
    ws2.auto_filter.ref = f"A1:F{len(tender_stats)+1}"

    wb.save(str(output_path))


# ═══════════════════════════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════════════════════════
def main():
    parser = argparse.ArgumentParser(description="Extract bidders from opening protocols")
    parser.add_argument("--downloads", default="downloads", help="Path to downloads folder")
    parser.add_argument("--inventory", default="inventory.xlsx", help="Path to inventory.xlsx")
    parser.add_argument("--output", default="bidder_extraction.xlsx", help="Output Excel path")
    parser.add_argument("--limit", type=int, default=0, help="Limit to N tenders (0=all)")
    parser.add_argument("--procurement-ids", default="", help="Comma-separated procurement IDs")
    parser.add_argument("--skip-llm", action="store_true", help="Skip LLM fallback for PDFs")
    args = parser.parse_args()

    downloads = Path(args.downloads)
    output = Path(args.output)

    # Determine which tender IDs to process
    if args.procurement_ids:
        target_ids = [int(x.strip()) for x in args.procurement_ids.split(",")]
    else:
        target_ids = sorted(
            int(d) for d in os.listdir(downloads)
            if d.isdigit() and (downloads / d).is_dir()
        )

    if args.limit > 0:
        target_ids = target_ids[:args.limit]

    print(f"Processing {len(target_ids)} tenders from {downloads}")
    print(f"Output: {output}")

    all_records = []
    tender_stats = {}
    last_api_call = 0

    for idx, pid in enumerate(target_ids):
        folder = downloads / str(pid)
        print(f"  [{idx+1}/{len(target_ids)}] {pid}...", end=" ", flush=True)

        records = extract_tender(folder, pid, skip_llm=args.skip_llm)
        records = _validate_records(records)
        all_records.extend(records)

        # Build stats
        valid = [r for r in records if not r.get("error")]
        errors = [r for r in records if r.get("error")]
        n_bidders = len(set(r.get("bidder_name", "") for r in valid))
        n_lots = len(set(r.get("lot", 1) for r in valid)) if valid else 0
        has_amounts = any(r.get("bid_amount") is not None for r in valid)
        source_types = set(r.get("source_type", "") for r in valid)
        proto_type = ", ".join(sorted(source_types)) if source_types else "none"

        if errors and not valid:
            status = "error"
        elif any(r.get("confidence") == "low" for r in records):
            status = "low_confidence"
        elif valid:
            status = "ok"
        else:
            status = "no_data"

        tender_stats[pid] = {
            "n_bidders": n_bidders, "n_lots": n_lots,
            "has_amounts": "yes" if has_amounts else "no",
            "protocol_type": proto_type, "status": status,
        }

        err_str = errors[0]["error"][:40] if errors else ""
        print(f"{n_bidders} bidders, {n_lots} lots, amounts={'yes' if has_amounts else 'no'}"
              + (f" [{err_str}]" if err_str else ""))

    write_output(all_records, output, tender_stats)

    # Summary
    statuses = pd.Series([s["status"] for s in tender_stats.values()])
    print(f"\nDone. {len(all_records)} records from {len(target_ids)} tenders.")
    print(f"  ok: {(statuses=='ok').sum()}")
    print(f"  low_confidence: {(statuses=='low_confidence').sum()}")
    print(f"  error: {(statuses=='error').sum()}")
    print(f"  no_data: {(statuses=='no_data').sum()}")
    print(f"Output: {output}")


if __name__ == "__main__":
    main()
